"""Multi-format document loader supporting PDF, DOCX, TXT, and Excel files."""

import os
from typing import List, Dict, Any
from pathlib import Path
import tempfile

from langchain.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


class DocumentLoader:
    """Load and process documents from multiple file formats."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document loader.

        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )

        # Try to import optional dependencies
        self._setup_optional_dependencies()

    def _setup_optional_dependencies(self):
        """Setup optional document processing libraries."""
        self.pdf_available = False
        self.docx_available = False
        self.excel_available = False

        try:
            import PyPDF2
            self.pdf_available = True
        except ImportError:
            pass

        try:
            from docx import Document as DocxDocument
            self.docx_available = True
            self.DocxDocument = DocxDocument
        except ImportError:
            pass

        try:
            import openpyxl
            self.excel_available = True
            import pandas as pd
            self.pd = pd
        except ImportError:
            pass

    def load_file(self, file_path: str, file_type: str = None) -> List[Document]:
        """
        Load a single file and return as documents.

        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, txt, excel). If None, inferred from extension.

        Returns:
            List of Document objects
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Infer file type if not provided
        if file_type is None:
            _, ext = os.path.splitext(file_path)
            file_type = ext.lower().lstrip('.')

        # Route to appropriate loader
        if file_type == 'pdf':
            return self._load_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return self._load_docx(file_path)
        elif file_type == 'txt':
            return self._load_txt(file_path)
        elif file_type in ['xlsx', 'xls', 'csv']:
            return self._load_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def load_directory(self, directory_path: str, file_types: List[str] = None) -> List[Document]:
        """
        Load all documents from a directory.

        Args:
            directory_path: Path to directory
            file_types: List of file types to load (e.g., ['pdf', 'txt']). If None, load all supported.

        Returns:
            List of all documents
        """
        documents = []
        supported_types = {'pdf', 'docx', 'doc', 'txt', 'xlsx', 'xls', 'csv'}

        if file_types is None:
            file_types = list(supported_types)

        for file_name in os.listdir(directory_path):
            file_path = os.path.join(directory_path, file_name)

            if os.path.isfile(file_path):
                _, ext = os.path.splitext(file_name)
                file_type = ext.lower().lstrip('.')

                if file_type in file_types and file_type in supported_types:
                    try:
                        docs = self.load_file(file_path, file_type)
                        documents.extend(docs)
                        print(f"Loaded {len(docs)} chunks from {file_name}")
                    except Exception as e:
                        print(f"Error loading {file_name}: {str(e)}")

        return documents

    def _load_txt(self, file_path: str) -> List[Document]:
        """Load text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        chunks = self.text_splitter.split_text(content)
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "file_type": "txt",
                    "chunk_id": i
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        return documents

    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load PDF file."""
        if not self.pdf_available:
            raise ImportError("PyPDF2 is required for PDF loading. Install with: pip install PyPDF2")

        import PyPDF2

        documents = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            full_text = ""
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                full_text += f"\n--- Page {page_num + 1} ---\n{text}\n"

        chunks = self.text_splitter.split_text(full_text)
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "file_type": "pdf",
                    "chunk_id": i
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        return documents

    def _load_docx(self, file_path: str) -> List[Document]:
        """Load DOCX file."""
        if not self.docx_available:
            raise ImportError("python-docx is required for DOCX loading. Install with: pip install python-docx")

        doc = self.DocxDocument(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        chunks = self.text_splitter.split_text(full_text)
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "file_type": "docx",
                    "chunk_id": i
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        return documents

    def _load_excel(self, file_path: str) -> List[Document]:
        """Load Excel or CSV file."""
        if not self.excel_available:
            raise ImportError("openpyxl and pandas required. Install with: pip install openpyxl pandas")

        try:
            # Use pandas to read Excel or CSV
            if file_path.endswith('.csv'):
                df = self.pd.read_csv(file_path)
            else:
                df = self.pd.read_excel(file_path)

            # Convert dataframe to text format
            full_text = "Data Table:\n"
            full_text += df.to_string()

            chunks = self.text_splitter.split_text(full_text)
            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "file_type": "excel",
                        "chunk_id": i,
                        "rows": len(df),
                        "columns": len(df.columns)
                    }
                )
                for i, chunk in enumerate(chunks)
            ]
            return documents
        except Exception as e:
            raise ValueError(f"Error reading Excel file: {str(e)}")

    def get_file_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """Get statistics about loaded documents."""
        return {
            "total_documents": len(documents),
            "total_characters": sum(len(doc.page_content) for doc in documents),
            "total_tokens_estimate": sum(len(doc.page_content.split()) for doc in documents),
            "sources": list(set(doc.metadata.get("source", "unknown") for doc in documents))
        }
